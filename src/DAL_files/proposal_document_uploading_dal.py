import os
import tempfile
import PyPDF2
import docx
from datetime import datetime
from typing import Dict, Optional, BinaryIO
import logging

# LangChain imports
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import HumanMessage

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentProcessingDAL:
    def __init__(self, api_key: str):
        """
        Initialize the Document Processing DAL using LangChain with Google Gemini

        Args:
            api_key (str): Google Gemini API key
        """
        os.environ["GOOGLE_API_KEY"] = api_key  # Required by LangChain
        self.model = ChatGoogleGenerativeAI(
            model="gemini-1.5-flash", temperature=0.2
        )
        logger.info("Document processor initialized using LangChain + Gemini 1.5 Flash")

    def extract_pdf_content(self, file_path: str) -> Dict:
        content = {
            'text': '',
            'pages': 0,
            'tables_detected': False,
            'extraction_method': 'PDF'
        }

        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                content['pages'] = len(pdf_reader.pages)

                full_text = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            full_text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                        continue

                content['text'] = full_text.strip()

                # Simple table detection
                table_indicators = ['|', '│', '┌', '┐', '└', '┘', '├', '┤', '┬', '┴', '┼']
                content['tables_detected'] = any(indicator in content['text'] for indicator in table_indicators)

            logger.info(f"PDF processed: {content['pages']} pages, {len(content['text'])} characters")
            return content

        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            raise Exception(f"Failed to extract PDF content: {str(e)}")

    def extract_docx_content(self, file_path: str) -> Dict:
        content = {
            'text': '',
            'paragraphs': 0,
            'tables_detected': False,
            'tables_content': [],
            'extraction_method': 'DOCX'
        }

        try:
            doc = docx.Document(file_path)

            paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            content['text'] = '\n\n'.join(paragraphs)
            content['paragraphs'] = len(paragraphs)

            if doc.tables:
                content['tables_detected'] = True
                tables_text = []

                for i, table in enumerate(doc.tables):
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_data.append(row_data)

                    table_text = f"\n--- Table {i + 1} ---\n"
                    table_text += "\n".join([" | ".join(row) for row in table_data])
                    tables_text.append(table_text)
                    content['tables_content'].append(table_data)

                content['text'] += '\n\n' + '\n'.join(tables_text)

            logger.info(f"DOCX processed: {content['paragraphs']} paragraphs, {len(doc.tables)} tables")
            return content

        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}")
            raise Exception(f"Failed to extract DOCX content: {str(e)}")

    def analyze_with_gemini(self, content: Dict) -> str:
        """
        Analyze document content using Gemini via LangChain

        Args:
            content (Dict): Document content and metadata

        Returns:
            str: AI-generated analysis summary
        """
        try:
            if not content.get('text') or len(content['text'].strip()) < 50:
                return "Document content is too short or empty for meaningful analysis."

            prompt = f"""
You are a senior sales analyst providing a comprehensive professional summary of this sales document. Write a detailed, structured summary that a C-level executive would expect to see. Also don't include any of the recomendations from your end, it must be fully based on the document.
[Important: in the summary resposne don't include you are giving a summary or the summary word]

Document Content:
{content['text']}
"""

            response = self.model.invoke([HumanMessage(content=prompt)])

            return response.content if hasattr(response, 'content') else str(response)

        except Exception as e:
            logger.error(f"Error analyzing with Gemini: {str(e)}")
            raise Exception(f"Failed to analyze document with AI: {str(e)}")

    def process_document(self, file_data: BinaryIO, filename: str) -> Dict:
        temp_file_path = None

        try:
            logger.info(f"Starting document processing: {filename}")

            if not filename.lower().endswith(('.pdf', '.docx', '.doc')):
                raise ValueError(f"Unsupported file type. Only PDF and DOCX files are supported. Got: {filename}")

            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as temp_file:
                temp_file_path = temp_file.name
                temp_file.write(file_data.read())
                temp_file.flush()
                file_data.seek(0)

            if filename.lower().endswith('.pdf'):
                content = self.extract_pdf_content(temp_file_path)
            else:
                content = self.extract_docx_content(temp_file_path)

            summary = self.analyze_with_gemini(content)

            result = {
                'filename': filename,
                'processing_timestamp': datetime.now().isoformat(),
                'file_size_kb': round(len(file_data.getvalue()) / 1024, 2) if hasattr(file_data, 'getvalue') else 0,
                'document_metadata': {
                    'extraction_method': content['extraction_method'],
                    'pages': content.get('pages', content.get('paragraphs', 0)),
                    'tables_detected': content.get('tables_detected', False),
                    'content_length': len(content['text']),
                    'has_content': len(content['text'].strip()) > 0
                },
                'summary': summary,
                'success': True,
                'error': None
            }

            logger.info(f"Document processing completed successfully: {filename}")
            return result

        except Exception as e:
            error_msg = str(e)
            logger.error(f"Document processing failed for {filename}: {error_msg}")
            return {
                'filename': filename,
                'processing_timestamp': datetime.now().isoformat(),
                'summary': None,
                'success': False,
                'error': error_msg,
                'document_metadata': None
            }

        finally:
            if temp_file_path and os.path.exists(temp_file_path):
                try:
                    os.unlink(temp_file_path)
                except Exception as e:
                    logger.warning(f"Failed to delete temporary file {temp_file_path}: {str(e)}")

    def validate_document(self, file_data: BinaryIO, filename: str) -> Dict:
        try:
            if not filename.lower().endswith(('.pdf', '.docx', '.doc')):
                return {
                    'valid': False,
                    'error': f"Unsupported file type. Only PDF and DOCX files are supported. Got: {filename}"
                }

            file_size = len(file_data.read())
            file_data.seek(0)

            if file_size > 10 * 1024 * 1024:
                return {
                    'valid': False,
                    'error': f"File size too large. Max size is 10MB. Got: {round(file_size / (1024*1024), 2)}MB"
                }

            if file_size == 0:
                return {
                    'valid': False,
                    'error': "File is empty"
                }

            return {
                'valid': True,
                'file_size_mb': round(file_size / (1024*1024), 2),
                'file_type': os.path.splitext(filename)[1].lower()
            }

        except Exception as e:
            return {
                'valid': False,
                'error': f"Validation error: {str(e)}"
            }
